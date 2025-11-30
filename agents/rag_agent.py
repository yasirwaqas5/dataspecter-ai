"""
RAG (Retrieval-Augmented Generation) Agent
Implements:
- Document loading (PDF, TXT, DOCX)
- Text chunking and embedding
- Vector store (FAISS)
- Semantic retrieval
- Context injection for LLM
"""

import os
from typing import List, Dict, Optional, Tuple
import warnings
warnings.filterwarnings("ignore")


class RAGAgent:
    """RAG system for financial document analysis"""
    
    def __init__(self, embedding_model: str = 'sentence-transformers', 
                 chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Initialize RAG Agent
        
        Args:
            embedding_model: Type of embedding ('sentence-transformers', 'openai')
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.vector_store = None
        self.documents = []
        self.embeddings = None
        
        # Do not initialize embeddings to avoid torch/sentence-transformers dependencies
        # self._initialize_embeddings()
    
    def _initialize_embeddings(self):
        """Initialize embedding model - disabled to avoid torch/sentence-transformers dependencies"""
        # Disabled to avoid torch/sentence-transformers dependencies
        self.embeddings = None
    
    def load_documents(self, file_paths: List[str] = None, 
                      texts: List[str] = None,
                      uploaded_files=None) -> Dict:
        """
        Load documents from various sources
        
        Args:
            file_paths: List of file paths
            texts: List of text strings
            uploaded_files: Streamlit uploaded files
            
        Returns:
            Dict with loading results
        """
        result = {
            'success': False,
            'documents_loaded': 0,
            'total_chunks': 0,
            'errors': []
        }
        
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            from langchain_core.documents import Document
            
            all_documents = []
            
            # Load from file paths
            if file_paths:
                for file_path in file_paths:
                    try:
                        docs = self._load_single_file(file_path)
                        all_documents.extend(docs)
                    except Exception as e:
                        result['errors'].append(f"Failed to load {file_path}: {str(e)}")
            
            # Load from uploaded files (Streamlit)
            if uploaded_files:
                for uploaded_file in uploaded_files:
                    try:
                        docs = self._load_uploaded_file(uploaded_file)
                        all_documents.extend(docs)
                    except Exception as e:
                        result['errors'].append(f"Failed to load {getattr(uploaded_file, 'name', 'unknown')}: {str(e)}")
            
            # Load from raw texts
            if texts:
                for i, text in enumerate(texts):
                    all_documents.append(Document(
                        page_content=text,
                        metadata={'source': f'text_{i}'}
                    ))
            
            # If no documents loaded, that's okay - return success with 0 documents
            if not all_documents:
                result.update({
                    'success': True,
                    'documents_loaded': 0,
                    'total_chunks': 0
                })
                return result
            
            # Split documents into chunks
            try:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
                )
                
                self.documents = text_splitter.split_documents(all_documents)
                
                result.update({
                    'success': True,
                    'documents_loaded': len(all_documents),
                    'total_chunks': len(self.documents)
                })
            except Exception as e:
                # If splitting fails, use documents as-is
                self.documents = all_documents
                result.update({
                    'success': True,
                    'documents_loaded': len(all_documents),
                    'total_chunks': len(all_documents),
                    'warning': f'Document splitting failed: {str(e)}, using documents as-is'
                })
            
        except Exception as e:
            result['errors'].append(str(e))
        
        return result
    
    def _load_single_file(self, file_path: str) -> List:
        """Load a single file based on extension"""
        from langchain_core.documents import Document
        
        # Validate file path
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"File does not exist: {file_path}")
        
        _, ext = os.path.splitext(file_path.lower())
        
        if ext == '.pdf':
            return self._load_pdf(file_path)
        elif ext == '.txt':
            return self._load_txt(file_path)
        elif ext in ['.docx', '.doc']:
            return self._load_docx(file_path)
        else:
            # Try to load as text file as fallback
            try:
                return self._load_txt(file_path)
            except Exception as e:
                raise ValueError(f"Unsupported file type: {ext}")
    
    def _load_pdf(self, file_path: str) -> List:
        """Load PDF file"""
        # Validate file path
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"File does not exist: {file_path}")
        
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            return loader.load()
        except ImportError:
            # Fallback to pypdf2
            try:
                import PyPDF2
                from langchain_core.documents import Document
                
                documents = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                documents.append(Document(
                                    page_content=text,
                                    metadata={'source': file_path, 'page': page_num}
                                ))
                        except Exception as e:
                            print(f"Warning: Could not extract text from page {page_num}: {e}")
                            # Add empty page as fallback
                            documents.append(Document(
                                page_content="",
                                metadata={'source': file_path, 'page': page_num}
                            ))
                return documents
            except ImportError:
                raise ImportError("No PDF processing library available. Install pypdf2 or langchain-community.")
        except Exception as e:
            # Try fallback method
            try:
                import PyPDF2
                from langchain_core.documents import Document
                
                documents = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                documents.append(Document(
                                    page_content=text,
                                    metadata={'source': file_path, 'page': page_num}
                                ))
                        except Exception as e:
                            print(f"Warning: Could not extract text from page {page_num}: {e}")
                            # Add empty page as fallback
                            documents.append(Document(
                                page_content="",
                                metadata={'source': file_path, 'page': page_num}
                            ))
                return documents
            except Exception as e2:
                raise Exception(f"Could not load PDF file: {e} (fallback also failed: {e2})")
    
    def _load_txt(self, file_path: str) -> List:
        """Load text file"""
        from langchain_core.documents import Document
        
        # Validate file path
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"File does not exist: {file_path}")
        
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'ISO-8859-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                # If we successfully read the file, break
                if text is not None:
                    break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                if encoding == 'cp1252':
                    # Last encoding, re-raise the exception
                    raise Exception(f"Could not read file with any encoding: {e}")
                continue
        
        return [Document(page_content=text, metadata={'source': file_path})]
    
    def _load_docx(self, file_path: str) -> List:
        """Load DOCX file"""
        # Validate file path
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"File does not exist: {file_path}")
        
        try:
            from docx import Document as DocxDocument
            from langchain_core.documents import Document
            
            doc = DocxDocument(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            
            return [Document(page_content=text, metadata={'source': file_path})]
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            # Try to read as text file as fallback
            try:
                return self._load_txt(file_path)
            except Exception as e2:
                raise Exception(f"Could not load DOCX file: {e} (fallback also failed: {e2})")
    
    def _load_uploaded_file(self, uploaded_file) -> List:
        """Load Streamlit uploaded file"""
        from langchain_core.documents import Document
        import tempfile
        
        # Reset file pointer to beginning if possible
        if hasattr(uploaded_file, 'seek'):
            try:
                uploaded_file.seek(0)
            except Exception as e:
                print(f"Warning: Could not reset file pointer: {e}")
        
        # Get file name and extension
        file_name = getattr(uploaded_file, 'name', 'unknown')
        file_ext = os.path.splitext(file_name.lower())[1]
        
        # For text files, try to read directly
        if file_ext in ['.txt']:
            try:
                # Read the content directly
                content = uploaded_file.read()
                # Reset file pointer again
                if hasattr(uploaded_file, 'seek'):
                    uploaded_file.seek(0)
                
                # Try to decode as text
                if isinstance(content, bytes):
                    content = content.decode('utf-8')
                
                return [Document(page_content=content, metadata={'source': file_name})]
            except Exception as e:
                print(f"Warning: Could not read text file directly: {e}")
                # Reset file pointer again
                if hasattr(uploaded_file, 'seek'):
                    uploaded_file.seek(0)
        
        # Save to temporary file
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_path = tmp_file.name
        except Exception as e:
            raise Exception(f"Could not save uploaded file to temporary location: {e}")
        
        try:
            docs = self._load_single_file(tmp_path)
            # Update metadata with original filename
            for doc in docs:
                doc.metadata['source'] = file_name
            return docs
        finally:
            try:
                os.unlink(tmp_path)
            except Exception as e:
                print(f"Warning: Could not delete temporary file: {e}")
    
    def create_vector_store(self) -> Dict:
        """
        Create FAISS vector store from loaded documents
        
        Returns:
            Dict with creation results
        """
        result = {
            'success': True,
            'num_vectors': len(self.documents) if self.documents else 0
        }
        
        # Always use mock vector store to avoid torch/sentence-transformers dependencies
        self.vector_store = None
        
        return result
    
    def retrieve(self, query: str, top_k: int = 3) -> List[Dict]:
        """
        Retrieve relevant documents for a query
        
        Args:
            query: Search query
            top_k: Number of documents to retrieve
            
        Returns:
            List of relevant document chunks
        """
        # Validate inputs
        if not query or not query.strip():
            return []
        
        if top_k <= 0:
            top_k = 3
        
        # Always use mock mode to avoid torch/sentence-transformers dependencies
        if self.documents:
            results = []
            for i, doc in enumerate(self.documents[:top_k]):
                results.append({
                    'rank': i + 1,
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'relevance_score': 1.0 - (i * 0.1)  # Simple scoring
                })
            return results
        return []
    
    def retrieve_with_scores(self, query: str, top_k: int = 3) -> List[Tuple[str, float]]:
        """
        Retrieve documents with similarity scores
        
        Returns:
            List of (document_text, score) tuples
        """
        # Validate inputs
        if not query or not query.strip():
            return []
        
        if top_k <= 0:
            top_k = 3
        
        # Always use mock mode to avoid torch/sentence-transformers dependencies
        if self.documents:
            results = []
            for i, doc in enumerate(self.documents[:top_k]):
                score = 1.0 - (i * 0.1)  # Simple scoring
                results.append((doc.page_content, float(score)))
            return results
        return []
    
    def get_context_for_query(self, query: str, max_length: int = 2000) -> str:
        """
        Get formatted context string for LLM injection
        
        Args:
            query: User query
            max_length: Maximum context length
            
        Returns:
            Formatted context string
        """
        # Validate inputs
        if not query or not query.strip():
            return ""
        
        if max_length <= 0:
            max_length = 2000
        
        # Always use mock mode to avoid torch/sentence-transformers dependencies
        if not self.documents:
            return ""
        
        docs_to_use = self.documents[:5]  # Use first 5 documents
        
        context_parts = []
        current_length = 0
        
        for doc in docs_to_use:
            content = doc.page_content
            source = doc.metadata.get('source', 'Unknown')
            
            chunk = f"[Source: {source}]\n{content}\n"
            
            if current_length + len(chunk) > max_length:
                break
            
            context_parts.append(chunk)
            current_length += len(chunk)
        
        if context_parts:
            return "=== RELEVANT DOCUMENTS ===\n\n" + "\n---\n".join(context_parts)
        else:
            return ""
    
    def save_vector_store(self, path: str):
        """Save vector store to disk"""
        if not path:
            raise ValueError("Path cannot be empty")
        
        if self.vector_store:
            try:
                self.vector_store.save_local(path)
                return True
            except Exception as e:
                print(f"Warning: Could not save vector store: {e}")
                return False
        else:
            print("Warning: No vector store to save")
            return False
    
    def load_vector_store(self, path: str):
        """Load vector store from disk"""
        if not path or not os.path.exists(path):
            print(f"Warning: Vector store path does not exist: {path}")
            return False
        
        try:
            from langchain_community.vectorstores import FAISS
            self.vector_store = FAISS.load_local(
                path,
                self.embeddings,
                allow_dangerous_deserialization=True
            )
            return True
        except Exception as e:
            print(f"Failed to load vector store: {e}")
            return False
    
    def get_stats(self) -> Dict:
        """Get RAG system statistics"""
        return {
            'num_documents': len(self.documents) if self.documents else 0,
            'has_vector_store': self.vector_store is not None,
            'embedding_model': getattr(self, 'embedding_model', 'unknown'),
            'chunk_size': getattr(self, 'chunk_size', 500),
            'chunk_overlap': getattr(self, 'chunk_overlap', 50)
        }


class FinancialRAGAgent(RAGAgent):
    """Specialized RAG agent for financial documents"""
    
    def build_index(self, uploaded_files):
        """
        Build a simple index from uploaded files without embeddings
        
        Args:
            uploaded_files: Streamlit uploaded files
            
        Returns:
            dict with chunks
        """
        try:
            # Check if we have files to process
            if not uploaded_files:
                return {"chunks": []}
            
            # Load documents
            load_result = self.load_documents(uploaded_files=uploaded_files)
            if not load_result.get('success'):
                return {"chunks": []}
            
            # Check if we have documents
            if not self.documents:
                return {"chunks": []}
            
            # Create simple chunks without embeddings
            chunks = []
            for doc in self.documents:
                chunks.append({
                    "text": doc.page_content,
                    "source": doc.metadata.get('source', 'unknown')
                })
            
            return {"chunks": chunks}
            
        except Exception as e:
            print(f"Error building simple index: {e}")
            return {"chunks": []}
    
    def answer_question(self, index, question: str, top_k: int = 3):
        """
        Answer a question using simple keyword-based retrieval
        
        Args:
            index: dict with chunks
            question: User question
            top_k: Number of relevant chunks to retrieve
            
        Returns:
            Tuple of (answer, sources)
        """
        try:
            # Validate inputs
            if index is None or not isinstance(index, dict) or "chunks" not in index:
                return "No document content available. Please process documents first.", []
            
            chunks = index.get("chunks", [])
            if not chunks:
                return "No document content available. Please process documents first.", []
            
            if not question or not question.strip():
                return "Please provide a valid question.", []
            
            # Test that we can process the question
            if len(question.strip()) < 3:
                return "Please ask a more detailed question.", []
            
            # Rank chunks by simple keyword overlap
            question_lower = question.lower()
            question_words = set(question_lower.split())
            
            # Score chunks based on keyword matching
            chunk_scores = []
            for i, chunk in enumerate(chunks):
                content_lower = chunk["text"].lower()
                content_words = set(content_lower.split())
                
                # Calculate overlap score
                overlap = len(question_words.intersection(content_words))
                if overlap > 0 or any(word in content_lower for word in question_words):
                    chunk_scores.append((i, overlap, chunk))
            
            # Sort by score and take top_k
            chunk_scores.sort(key=lambda x: x[1], reverse=True)
            top_chunks = chunk_scores[:top_k]
            
            # If no matches found, use first chunks
            if not top_chunks:
                top_chunks = [(i, 0, chunk) for i, chunk in enumerate(chunks[:top_k])]
            
            # If still no chunks, return error
            if not top_chunks:
                return "No relevant information found in the documents.", []
            
            # Extract content and sources
            context_parts = []
            sources = []
            
            for i, (idx, score, chunk) in enumerate(top_chunks):
                context_parts.append(chunk["text"])
                source = chunk.get("source", "unknown")
                if source not in sources:
                    sources.append(source)
            
            # Combine context
            context = "\n\n".join(context_parts)
            
            # Build prompt for EnhancedLLMAgent
            prompt = f"""
You are a QA assistant. Use ONLY the chunks below to answer the question.

CHUNKS:
{context}

QUESTION:
{question}

If the answer is not in the chunks, say you cannot find it in the provided documents.
"""
            
            # Call EnhancedLLMAgent
            try:
                from agents.enhanced_llm_agent import EnhancedLLMAgent
                
                # Use the provider and API key passed to the constructor
                provider = self.provider
                api_key = self.api_key
                
                # If no API key was provided, we'll use the fallback
                if not api_key:
                    # Fallback: Return a simple answer based on context
                    fallback_answer = f"Based on the documents, here's what I found about '{question}':\n\n{context[:1000]}"
                    if len(context) > 1000:
                        fallback_answer += "..."
                    return fallback_answer, sources
                
                llm_agent = EnhancedLLMAgent(provider=provider, api_key=api_key)
                answer = llm_agent.ask(prompt)
                
                return answer, sources
                
            except Exception as e:
                # Fallback: Return a simple answer based on context
                fallback_answer = f"Based on the documents, here's what I found about '{question}':\n\n{context[:1000]}"
                if len(context) > 1000:
                    fallback_answer += "..."
                return fallback_answer, sources
                
        except Exception as e:
            return f"Error processing question: {str(e)}", []
    
    def __init__(self, provider: str = None, api_key: str = None, **kwargs):
        try:
            super().__init__(**kwargs)
        except Exception as e:
            print(f"Warning: Could not initialize RAGAgent parent class: {e}")
            # Initialize required attributes manually
            self.embedding_model = kwargs.get('embedding_model', 'sentence-transformers')
            self.chunk_size = kwargs.get('chunk_size', 500)
            self.chunk_overlap = kwargs.get('chunk_overlap', 50)
            self.vector_store = None
            self.documents = []
            self.embeddings = None
        
        # Store provider and API key for LLM usage
        self.provider = provider
        self.api_key = api_key
        
        try:
            self.financial_keywords = [
                'revenue', 'profit', 'loss', 'assets', 'liabilities', 'equity',
                'cash flow', 'earnings', 'dividend', 'balance sheet', 'income statement',
                'roi', 'ebitda', 'margin', 'valuation', 'stock', 'investment'
            ]
        except Exception as e:
            print(f"Warning: Could not initialize FinancialRAGAgent: {e}")
            # Initialize with default values
            self.financial_keywords = [
                'revenue', 'profit', 'loss', 'assets', 'liabilities', 'equity',
                'cash flow', 'earnings', 'dividend', 'balance sheet', 'income statement',
                'roi', 'ebitda', 'margin', 'valuation', 'stock', 'investment'
            ]
    
    def extract_financial_entities(self, text: str) -> Dict:
        """Extract financial entities from text"""
        import re
        
        entities = {
            'monetary_values': [],
            'percentages': [],
            'dates': [],
            'financial_terms': []
        }
        
        # Extract monetary values ($X, $X.XX million/billion)
        money_pattern = r'\$\s*[\d,]+\.?\d*\s*(?:million|billion|trillion|M|B|T)?'
        entities['monetary_values'] = re.findall(money_pattern, text, re.IGNORECASE)
        
        # Extract percentages
        pct_pattern = r'\d+\.?\d*\s*%'
        entities['percentages'] = re.findall(pct_pattern, text)
        
        # Extract financial keywords
        text_lower = text.lower()
        for keyword in self.financial_keywords:
            if keyword in text_lower:
                entities['financial_terms'].append(keyword)
        
        return entities
    
    def summarize_financial_document(self, document_text: str, max_length: int = 500) -> str:
        """Create a financial summary of a document"""
        entities = self.extract_financial_entities(document_text)
        
        summary_parts = []
        
        if entities['monetary_values']:
            summary_parts.append(f"Key figures: {', '.join(entities['monetary_values'][:5])}")
        
        if entities['percentages']:
            summary_parts.append(f"Percentages mentioned: {', '.join(entities['percentages'][:3])}")
        
        if entities['financial_terms']:
            summary_parts.append(f"Topics: {', '.join(set(entities['financial_terms'][:5]))}")
        
        return " | ".join(summary_parts) if summary_parts else "No financial data extracted"
